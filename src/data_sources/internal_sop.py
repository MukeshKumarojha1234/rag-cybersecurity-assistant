"""Internal SOP data source (node 1).

Loads the organization's own incident-response / vulnerability-management
procedures from a local directory so they can be retrieved and cited
alongside public threat intel.
"""
from __future__ import annotations

import logging
from datetime import datetime, timezone
from pathlib import Path

from . import config
from .base import DataSource
from .models import RawDocument

logger = logging.getLogger(__name__)

_TEXT_EXTENSIONS = {".md", ".txt", ".log", ".csv"}


class InternalSopSource(DataSource):
    source_type = "internal_sop"

    def __init__(self, sop_dir: Path | str | None = None):
        self.sop_dir = Path(sop_dir or config.INTERNAL_SOP_DIR)

    def fetch(self) -> list[RawDocument]:
        if not self.sop_dir.exists():
            logger.warning("Internal SOP directory %s does not exist; returning no documents.", self.sop_dir)
            return []

        documents = []
        for path in sorted(self.sop_dir.rglob("*")):
            if not path.is_file() or path.stem.upper() == "README":
                continue
            doc = self._load_file(path)
            if doc:
                documents.append(doc)
        return documents

    def _load_file(self, path: Path) -> RawDocument | None:
        suffix = path.suffix.lower()
        try:
            if suffix in _TEXT_EXTENSIONS:
                content = path.read_text(encoding="utf-8", errors="replace")
            elif suffix == ".pdf":
                content = self._read_pdf(path)
            elif suffix == ".docx":
                content = self._read_docx(path)
            else:
                logger.debug("Skipping unsupported SOP file type: %s", path)
                return None
        except Exception:
            logger.exception("Failed to read SOP file %s", path)
            return None

        if not content or not content.strip():
            return None

        stat = path.stat()
        return RawDocument(
            id=path.stem,
            source_type="internal_sop",
            title=path.stem.replace("_", " ").replace("-", " ").title(),
            content=content.strip(),
            url=str(path),
            published_at=datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc),
            metadata={"file_name": path.name, "relative_path": str(path.relative_to(self.sop_dir))},
        )

    @staticmethod
    def _read_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("pypdf not installed; skipping %s (pip install pypdf).", path)
            return ""
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            import docx
        except ImportError:
            logger.warning("python-docx not installed; skipping %s (pip install python-docx).", path)
            return ""
        document = docx.Document(str(path))
        return "\n".join(p.text for p in document.paragraphs)
